import os
from dotenv import load_dotenv
from docx import Document
import docx
from openai import OpenAI
import docx
import pandas as pd
import streamlit as st


load_dotenv()
model = os.getenv('LLM_MODEL', 'gpt-4o')


def GetRunningMemory(sessionState):
    """This functions will recuperate running memory from sessions state messages
    
    sessionstate:; st.sessionstate (streamlit running memory)
    """
    full_message = "\n\n -".join(message['content'] for message in sessionState['messages'])

    return full_message


def ProcessSummary(RunningMemory):
    """ 
    This function will transform the running memory (interview summary) into a comprehensive process description in Natural Language. 
    This description will then be used to generate feedback and identify missing elements or inconsistency. 

    """
    client = OpenAI()

    response = client.chat.completions.create(
        messages= [{"role": "user", "content": "Here is a conversation talking about a process: \n Conversation: \n "+  str (RunningMemory)},
                   {"role": "user", "content": "Provide a detailed view on the process flow highlighting BPMN elements needed to model a process such as: Actors, Tasks, Systems, Gateways..."},
                   {"role": "user", "content": "Outputs must be only the Process description following structure template, no additionnal comments"}, 
                   {"role": "user", "content": "Outputs must contains only standardized caracter, not bold caracter"}],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0, 
        model= model
        )
    response = response.choices[0].message.content
    return response

def FlowSummary(RunningMemory):
    """ 
    This function will transform the running memory (interview summary) into a comprehensive process description in Natural Language. 
    This description will then be used to generate feedback and identify missing elements or inconsistency. 

    """
    client = OpenAI()

    response = client.chat.completions.create(
        messages= [{"role": "user", "content": "Here is a conversation talking about a process: \n Conversation: \n "+  str (RunningMemory)},
                   {"role": "user", "content": "Summarize the very high level process flow using a structured natural text."},
                   {"role": "user", "content": "Outputs must be only the Process description easy to read, no additionnal comments"}, 
                   {"role": "user", "content": "Outputs must contains only standardized caracter, not bold caracter"}],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0, #Random (1) vs deterministic (0)
        model= model
        )
    response = response.choices[0].message.content
    return response

def graphCreation(information, mermaid): # MERMAID
    """"
    This functions will generate a graph that need to be print in streamlit chatbot.
    """
    client = OpenAI()
    response = client.chat.completions.create(
        messages= [{"role": "user", "content": "Your goal is to create or modify the following mermaid flowchart diagram. \n"},
                   {"role": "user", "content": "Only Flowchart TD diagram \n"},
                     {"role": "user", "content":  "Generate a simple flowchart diagram based on these informations: " + str(information)},
                     {"role": "user", "content":  "The flowchart diagram must contain ONLY provided information. No more informations " },
                        {"role": "user", "content": "If a mermaid diagram is provided in the description below. Modify the existing mermaid diagram and integrate information from the last answer"},
                        {"role": "user", "content":  "Here is the mermaid diagram to modify (if already existing): \n" + str(mermaid) },
                        {"role": "user", "content":  "the mermaid diagram must fist with container width define in streamlit: width = 600 \n" + str(mermaid) },
                      {"role": "user", "content": "Do not provide any additional information or notes, ONLY markdown."}, 
                       ],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0.5, 
        model= model
        )
    response = response.choices[0].message.content

    return response[10: -4]



def read_bpmn_files(folder_path):
    # List to hold the contents of each .bpmn file
    bpmn_contents = []

    # Iterate over all files in the specified folder
    for filename in os.listdir(folder_path):
        # Check if the file has a .bpmn extension
        if filename.endswith('.bpmn'):
            # Construct the full file path
            file_path = os.path.join(folder_path, filename)
            # Open and read the file
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # Append the content to the list
                bpmn_contents.append(content)
    
    return bpmn_contents


 # Function to read and display the contents of the file
def read_bpmn(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        return content

# Function to write new content to the file
def write_bpmn(content):
    new_file_path = "BPMN in progress.bpmn"
    with open(new_file_path, 'w', encoding='utf-8') as file:
        file.write(content)
        print("Content written to BPMN file.")


def supprimer_contenu_bpmn(path):
    file_path = path    
    # Ouvrir le fichier en mode 'w' pour supprimer tout le contenu
    with open(file_path, 'w', encoding='utf-8') as file:
        pass  # Rien n'est écrit, ce qui vide le fichier

    #print("Le contenu du fichier BPMN a été entièrement supprimé.")

# Function to read a file
def read_file(file_path):
    """
    This function will allow to read provided files (.docx).

    Parameters:
    -----------
    file_path: Path to retrieve document
    """
    file = open(file_path,"r")

    text= file.readlines()
    #print(text)
    return text


def read_docx(path):
    document = Document(path)
    memoryupd = ""
    for p in document.paragraphs:
        #print(p.text)
        memoryupd += p.text
        memoryupd +='\n'
    return memoryupd

def read_Q(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n\n'.join(full_text)

def remove_first_paragraph(UpdatedQuestionnairePath):
    # Read the document
    text = read_Q(UpdatedQuestionnairePath)

    # Split the text into paragraphs based on '\n\n'
    paragraphs = text.split('\n\n')

    # Remove the first paragraph if it exists
    if paragraphs:
        paragraphs = paragraphs[1:]

    # Create a new document and add the remaining paragraphs
    new_doc = Document()
    for paragraph in paragraphs:
        new_doc.add_paragraph(paragraph)

    # Save the modified document
    new_doc.save('\\Feed\\UpdatedQuestionnaire.docx')

def remove_first_lign(filepath):
    """
    This functions will allow to remove the first lign of the excel sheet when the question is already asked.

    Parameters:
    -----------
    excel_file: File that contains the questions to ask to the stakeholder.

    Output:
    -------
    excel_file: excel file without the first lign 


    """
    Questiondf = pd.read_excel(filepath)
    Questiondf = Questiondf[1:]
    Questiondf.reset_index()
    Questiondf.to_excel("Feed\\Questionnaire.xlsx")

    
#Display graph (Mermaid)

import base64
from IPython.display import Image, display

def mm_ink(graphbytes):
  """Given a bytes object holding a Mermaid-format graph, return a URL that will generate the image."""
  base64_bytes = base64.b64encode(graphbytes)
  base64_string = base64_bytes.decode("ascii")
  return "https://mermaid.ink/img/" + base64_string

def mm_display(graphbytes):
  """Given a bytes object holding a Mermaid-format graph, display it."""
  display(Image(url=mm_ink(graphbytes)))

def mm(graph):
  """Given a string containing a Mermaid-format graph, display it."""
  graphbytes = graph.encode("ascii")
  mm_display(graphbytes)


import hashlib
import random

def generate_hash_with_random_digits(input_str):
    # Generate a hash from the input string
    hash_obj = hashlib.sha256(input_str.encode())
    hash_code = hash_obj.hexdigest()
    
    # Append 10 random digits to the hash
    random_digits = ''.join(str(random.randint(0, 9)) for _ in range(10))
    complete_code = hash_code + random_digits
    
    return complete_code
