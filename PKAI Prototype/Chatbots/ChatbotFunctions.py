# Dictionary based on your provided structure
import os
from dotenv import load_dotenv
from openai import OpenAI
import streamlit as st
import json
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()
model = os.getenv('LLM_MODEL', 'gpt-4o')

# Function to save the updated JSON data back to the file
def save_json(filepath, data):
    with open(filepath, 'w') as file:
        json.dump(data, file, indent=4)
    print("Changes have been saved.")



def CheckAndFollowUp(prompt, GranularityScore):
    """
    This function is responsible to coordinate all the checks and select the next (follow-up or not) question
    """

    check(question = st.session_state.messages[-1]["content"], prompt = prompt, GranularityScore = GranularityScore)

        #Question to close the interview
    if len(st.session_state["CANVA"]["Questions"]) == 0:
        st.session_state["CANVA"]["Questions"].insert(0, "The interview is now over. Would you like to add anything? If not, press “Stop” to end the interview.")
        question = st.session_state["CANVA"]["Questions"].pop(0)
        st.session_state["CANVA"]["Expected Outputs"].insert(0, "None")

    # Check if the category exists and is not empty
    if len(st.session_state["CANVA"]["Questions"]) > 0:
        # Remove the first question from the category
        first_question = st.session_state["CANVA"]["Questions"].pop(0)
        question = first_question
        prompt = prompt

    return prompt, question
        

# Quality check (Dynamic Chatbot):
def check(question, prompt, GranularityScore):
    if len(st.session_state["CANVA"]["Expected Outputs"])>0:
        Expected_Outputs= st.session_state["CANVA"]["Expected Outputs"].pop(0)
    
    else:
        Expected_Outputs = "None"

    if checkIDK(question, prompt) ==  False:
        print (st.session_state["CANVA"]["Expected Outputs"][0])
        if st.session_state["CANVA"]["Expected Outputs"][0]!= "None":
            print ("check with expected outputs")
            if check_expected(expected= Expected_Outputs, question = question, prompt = prompt) < GranularityScore:
                follow_up_question = GenFollowUp(question, prompt, Expected_Outputs)     
                        
                # Append the follow-up question
                st.session_state["CANVA"]["Questions"].insert(0,follow_up_question)
                st.session_state["CANVA"]["Expected Outputs"].insert(0, "None")

        else:
            print ("check with provided information")
            checkInfo(question, prompt)

# Function to check information (dummy function)

def check_expected(expected, question, prompt):
    """
    This function checks if the answer provides sufficient information about the expected output element.
    
    Parameters:
    Question: Previous Questions
    Expected: The expected output element to focus on.
    answer: The provided answer to check.

    Returns:
    IDK: A boolean indicating whether the answer provides sufficient information (True) or not (False).
    """
    client = OpenAI()
    response = client.chat.completions.create(
        messages=st.session_state["messages"][-50:] + [
                  {"role": "user", "content": f"Here is the answer to a question. \n Question: {question} \n Answer: {prompt}. \n\n Can you check if the answer provides information about these elements: {expected}?"},
                  {"role": "user", "content": "Provide a short answer and give a score from 1 to 10. Provide a bad score (less than 5) if there is missing elements or if the quality of the answer is really poor. You must be very demanding"}
        ],
        max_tokens=1,
        n=1,
        stop=None,
        temperature=0.1,  # Random (1) vs deterministic (0)
        model=model
    )
    response_content = response.choices[0].message.content
    try:
        # Assume the score is provided as a number in the response
        score = int(''.join(filter(str.isdigit, response_content)))
    except ValueError:
        score = 0  # Default to 0 if no valid score is found

    #print ("Expected output check score (above 4 on ten means not enough details concerning expected outputs):" + str(score))
    print ("Expected output check score (below 4 on ten means not enough details provided):" + str(score))
    return score


# Function to generate follow-up questions (dummy function)
def GenFollowUp(question, prompt, Expected_Outputs):
    client = OpenAI()
    response = client.chat.completions.create(
        messages= st.session_state["messages"][-50:] + [
                   {"role": "assistant", "content": " Here is the question asked: " + str(question)},
                    {"role": "user", "content": "Here is the answer to the question: " + str(prompt) +". \n The answer does not provide sufficient information about some of these elements:  " + str (Expected_Outputs) +". Please, generate a follow-up question in order to get insights about the elements that are not discussed. Ask a process oriented question understandable for novice stakeholders."},
                    {"role": "user", "content":"The output must be the questions without any comments!"}, 
                    {"role": "user", "content":"No more than 2 questions in the output"}],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0.1, 
        model= model
        )
    Question = response.choices[0].message.content

    return Question



def checkIDK(question, answer):
    """This function will check if the end-users what to express that he/she does not know about the topic asked. You Should ask to someone esle !"""

    client = OpenAI()
    response = client.chat.completions.create(
        messages= st.session_state["messages"][-50:] + [
                    {"role": "user", "content": " Here is the question asked: " + str(question) + "\n. Here is the answer to the question: " + str(answer) +"."},
                   {"role": "user", "content": "  Check if the answer wants to express that the interviewed people have not the information"},
                    {"role": "user", "content": " Some clue should be expression such as: 'i don't know', 'ask to another person', 'i'm not responsible of that', 'i'm not sure', 'nothing more to say', 'Not relevant', 'not my scope', 'not defined'..."},
                    {"role": "user", "content": " if you diagnosed this in the answer, the output must be '1', else the output must contain '0'"}
                    ],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0.1, 
        model= model
        )
    IDK = False
    idk_response = response.choices[0].message.content
    if "1" in idk_response:
        IDK = True
    print ("\n ----------------------------------- \n IDK: (false = not detected) "  + str(IDK))
    return IDK



def checkInfo(question, prompt):
    client = OpenAI()
    response = client.chat.completions.create(
        messages= st.session_state["messages"][-50:] + [
                   {"role": "assistant", "content": " Here is the question asked: " + str(question)},
                    {"role": "user", "content":"Here is the answer to the question: " + str(prompt)},
                    {"role": "system", "content": "If the answer seems relevant to the topic under discussion and provides interesting information, score = 0, if the answer seems not logical according to the question, score = 1. Returns only the score. You must be very demanding."},
                    {"role": "user", "content":" If it is the first question and answer, score = 0 except if the user is asking practical use of the tools " + str(prompt)}],
        max_tokens=None,
        n=1,
        stop=None,
        temperature=0.1, 
        model= model
        )
    response = response.choices[0].message.content
    print("response: " + str(response))
    try:
        # Assume the score is provided as a number in the response
        score = int(''.join(filter(str.isdigit, response)))
    except ValueError:
        score = 0  # Default to 0 if no valid score is found

    if score == 1:
        response = client.chat.completions.create(
        messages= st.session_state["messages"][-50:] + [
                    {"role": "assistant", "content": " Here is the question asked: " + str(question) },
                   {"role": "user", "content":"Here is the answer to the question: " + str(prompt)},
                   {"role": "system", "content":"It seems that the answer doesn't provide all required information, generate a follow up question or reformulate the questions. Output must be only the question without any other comments."},
                   {"role": "user", "content":"No more than 1 question in the output"}],
        max_tokens= None,
        n=1,
        stop=None,
        temperature=0.5, 
        model= model
        )
        response = response.choices[0].message.content

        st.session_state["CANVA"]["Questions"].insert(0, response)
        st.session_state["CANVA"]["Expected Outputs"].insert(0, "None")

    print ("Check info score: (0 = infos ok, 1 = info ko ) " + str(score))


def Reporting(ProcessFlow, ProcessMemory, Memory, feedback):
    # Create or open the document
    
        # Define the file path
    file_path = "Outputs.docx"


    doc = Document()

    # Insert Title
    title = doc.add_heading(level=1)
    title_run = title.add_run("Process Documentation")
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading(level=2)
    subtitle_run = subtitle.add_run("Process Description")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Add introductory sentence
    doc.add_paragraph("Here is the high level model captured by the interviews.")
    # Insert Image
    doc.add_picture("Outputs/bpmn (1).png", width=Inches(6.0))  # Adjust width as necessary


    # Add additional description text
    doc.add_paragraph("This section provides a high level description of the process flow...")
    doc.add_paragraph(ProcessFlow)

    doc.add_paragraph() # ------------------------------------------------------------------------------------

    
    subtitle = doc.add_heading(level=2)
    subtitle_run = subtitle.add_run("Process Detailed Description")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("Here is the more detailed model captured by the interviews.")
    # Insert Image
    doc.add_picture("Outputs/bpmn.png", width=Inches(6.0)) 

    # Add Space (by adding a blank paragraph)
    doc.add_paragraph()

        # Insert Subtitle
    subtitle = doc.add_heading(level=2)
    subtitle_run = subtitle.add_run("Process Structured Documentation")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add additional description text
    doc.add_paragraph("This section provides a detailed description of the documented process...")
    doc.add_paragraph(ProcessMemory)
    doc.add_paragraph() # ------------------------------------------------------------------------------------


        # Insert Subtitle
    subtitle = doc.add_heading(level=2)
    subtitle_run = subtitle.add_run("Conversation with Domain Experts")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add additional description text
    doc.add_paragraph("This section provides a detailed description of the documented process...")
    doc.add_paragraph(Memory)

    message = " "
    if feedback:
        for e in feedback:
            message += "\n\n"
            message += e

            # Insert Subtitle
    subtitle = doc.add_heading(level=2)
    subtitle_run = subtitle.add_run("Comments on the Mermaid Model")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("These are the feedback on the mermaid model")
    doc.add_paragraph(message)
    

    # Save the document
    doc.save("Outputs.docx")



