import os
import json
import fitz  # PyMuPDF for PDF files
import docx  # For DOCX files
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.schema import Document
from openai import OpenAI

# Load environment variables from .env file
load_dotenv()
client = OpenAI()

# Retrieve API key from environment variables
api_key = os.getenv()
model = os.getenv('LLM_MODEL', 'gpt-4o')  # Default to 'gpt-4'

# Set up ChatOpenAI with updated configurations
llm = ChatOpenAI(api_key=api_key, model=model)

# Function to read text from a PDF file
def read_pdf(file_path):
    text = ""
    with fitz.open(file_path) as pdf_document:
        for page in pdf_document:
            text += page.get_text()
    return text

# Function to read text from a DOCX file
def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Load all documents from a given folder
def load_documents_from_folder(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith('.pdf'):
            print(f"Reading {file_path}...")
            text = read_pdf(file_path)
        elif filename.endswith('.docx'):
            print(f"Reading {file_path}...")
            text = read_docx(file_path)
        else:
            continue
        
        segments = [text[i:i + 1000] for i in range(0, len(text), 1000)]
        for segment in segments:
            documents.append(Document(page_content=segment, metadata={"filename": filename}))
    return documents

# Load existing questions from JSON
def load_existing_questions(json_path):
    with open(json_path, 'r') as file:
        existing_questions = json.load(file)
    return existing_questions

def generate_questions(input_process):
    try:
        # API call to generate questions and expected outputs
        response = client.chat.completions.create(
            model="gpt-4o-2024-08-06",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a business process analyst tasked with developing an interview guide to discover processes."},
                {"role": "user", "content":
                    f"Based on the following process structure, generate tailored interview questions:\n\n{input_process}\n\n"
                    "Please structure the output according to the provided schema:\n\n"}
            ],
            functions=[{
                "name": "process_schema",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "Questions": {
                            "type": "object",
                            "properties": {
                                "Before": {"type": "array", "description": "Questions to consider before starting the process.", "items": {"type": "string"}},
                                "StartEvent": {"type": "array", "description": "Questions related to the initiation of the process.", "items": {"type": "string"}},
                                "ProcessCore": {"type": "array", "description": "Questions that describe the main activities of the process.", "items": {"type": "string"}},
                                "EndEvent": {"type": "array", "description": "Questions that determine the completion of the process.", "items": {"type": "string"}},
                                "After": {"type": "array", "description": "Questions for discussion after the process has been completed.", "items": {"type": "string"}},
                                "Closing": {"type": "array", "description": "Final prompt to close the discussion.", "items": {"type": "string"}}
                            },
                            "required": ["Before", "StartEvent", "Process Core", "EndEvent", "After", "Closing"],
                            "additionalProperties": False
                        },
                        "ExpectedOutputs": {
                            "type": "object",
                            "properties": {
                                "Before": {"type": "array", "description": "Expected outputs and conditions before to start the analysis of the process.", "items": {"type": "string"}},
                                "StartEvent": {"type": "array", "description": "Expected outputs and conditions for the start event.", "items": {"type": "string"}},
                                "ProcessCore": {"type": "array", "description": "Expected outputs detailing the core activities of the process.", "items": {"type": "string"}},
                                "EndEvent": {"type": "array", "description": "Expected outputs to indicate process completion.", "items": {"type": "string"}},
                                "After": {"type": "array", "description": "Expected outputs and conditions after the process analysis.", "items": {"type": "string"}},
                                "Closing": {"type": "array", "description": "Expected input Modelling ", "items": {"type": "string"}}
                            },
                            "required": ["Before","StartEvent", "ProcessCore", "EndEvent", "After", "Closing"],
                            "additionalProperties": False
                        }
                    },
                    "required": ["Questions", "ExpectedOutputs"],
                    "additionalProperties": False
                }
            }]
        )
        
        # Extract JSON from the function_call arguments
        if response.choices[0].message.function_call:
            arguments = response.choices[0].message.function_call.arguments
            generated_data = json.loads(arguments)
            
            # Ensure every question has a corresponding Expected Output entry
            for section in ["Before","StartEvent", "ProcessCore", "EndEvent", "After", "Closing"]:
                questions = generated_data["Questions"].get(section, [])
                expected_outputs = generated_data["ExpectedOutputs"].get(section, [])
                
                # If expected outputs are fewer than questions, extend with empty strings
                if len(expected_outputs) < len(questions):
                    expected_outputs.extend([""] * (len(questions) - len(expected_outputs)))
                
                # Update the expected outputs in generated data
                generated_data["ExpectedOutputs"][section] = expected_outputs
            
            return generated_data
        else:
            print("The response does not contain function call arguments.")
            return None
    
    except Exception as e:
        print(f"An error occurred while generating questions: {e}")
        return None



# Configure the chatbot for document analysis
def configure_chatbot(folder_path):
    # Load documents
    documents = load_documents_from_folder(folder_path)

    # Create embeddings for your documents
    embeddings = OpenAIEmbeddings(api_key=api_key)
    vectorstore = Chroma.from_documents(documents, embeddings)

    # Set up the QA retrieval chain
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever()
    )
    return qa_chain

# Main function to run the application
def main():
    folder_path = input("Enter the path to the folder containing the documents: ")
    json_path = input("Enter the path to the JSON file with existing questions: ")
    
    qa_chain = configure_chatbot(folder_path)
    
    # Load existing questions
    existing_questions = load_existing_questions(json_path)

    # Define process structure (hard-coded or loaded)
    input_process = """Define your process structure here"""

    # Generate interview questions based on the input process and existing questions
    generated_questions = generate_questions(input_process, existing_questions)

    # Specify the output path for the new JSON file
    output_path = "generated_questions_output.json"

    # Check if generated_questions is not None and save it to a JSON file
    if generated_questions:
        with open(output_path, 'w') as output_file:
            json.dump(generated_questions, output_file, indent=2)
        print(f"Generated Questions JSON saved to {output_path}")
    else:
        print("Failed to generate questions.")

if __name__ == "__main__":
    main()
