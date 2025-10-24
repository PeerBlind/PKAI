import os
from PyPDF2 import PdfReader
import docx  # Pour les fichiers DOCX
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI  # Utilisation de ChatOpenAI pour les modèles de chat
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.schema import Document

# Chargez les variables d'environnement depuis le fichier .env
load_dotenv()

# Récupérez la clé API depuis les variables d'environnement
api_key = os.getenv()
model = os.getenv('LLM_MODEL', 'gpt-4o')  # Utilisez 'gpt-4o' par défaut
llm = ChatOpenAI(api_key=api_key, model=model)  # Utilisez ChatOpenAI

# Fonction pour lire le texte d'un fichier PDF
def read_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    # Extract text from each page
    for page in reader.pages:
        print(page.extract_text())
        text += page.extract_text()
    return text

# Fonction pour lire le texte d'un fichier DOCX
def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Chargez tous les documents d'un dossier donné
def load_documents_from_folder(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith('.pdf'):
            print(f"Reading {file_path}...")
            text = read_pdf(file_path)
        elif filename.endswith('.docx'):
            print(f"Reading {file_path}...")
            text = read_docx(file_path)
        else:
            continue
        
        # Segmentez le texte en morceaux plus petits si nécessaire
        segments = [text[i:i + 1000] for i in range(0, len(text), 1000)]
        for segment in segments:
            documents.append(Document(page_content=segment, metadata={"filename": filename}))
    return documents

# Fonction pour configurer le chatbot
def configure_chatbot(folder_path):
    # Chargez les documents
    documents = load_documents_from_folder(folder_path)

    # Créez un vecteur pour vos documents
    embeddings = OpenAIEmbeddings(api_key=api_key)
    vectorstore = Chroma.from_documents(documents, embeddings)

    # Configurez le système de récupération de QA
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever()
    )
    return qa_chain

# Fonction pour poser des questions
def ask_question(qa_chain, question: str) -> str:
    try:
        # Passer la clé 'query' avec la question
        response = qa_chain.invoke({"query": question})  
        print("Raw response:", response)  

        if 'result' in response:
            return response['result']  
        else:
            return "La réponse ne contient pas la clé 'result'. Vérifiez les paramètres de la chaîne."
    except Exception as e:
        print(f"An error occurred: {e}")  
        return str(e)

# Fonction principale
def main(folder_path):
    qa_chain = configure_chatbot(folder_path)

    while True:
        question = input("Posez votre question (ou tapez 'exit' pour quitter): ")
        if question.lower() == 'exit':
            break
        response = ask_question(qa_chain, question)
        print("Réponse:", response)

if __name__ == "__main__":
    folder_path = "Documentation"
    main(folder_path)
